import os
import sys
import tkinter as tk
from tkinter import ttk
from tkinter import filedialog
from docx import Document
from docx.shared import Pt


# --- Define global variables for muliple experiments ---
experiments = {}  # Dictionary to store all experiment data
current_experiment = None  # Variable to store the name of the current experiment
biologicals_list = []  # List to store biologicals for the current experiment
supplies_list = []  # List to store supplies for the current experiment
media_list = []  # List to store media for the current experiment
chemicals_list = []  # List to store chemicals for the current experiment

# --- Ensure the Libraries folder is in sys.path ---
script_dir = os.path.dirname(os.path.abspath(__file__))
libraries_path = os.path.join(script_dir, "..", "Libraries")
if libraries_path not in sys.path:
    sys.path.append(libraries_path)

# --- Import custom libraries ---
import bacteria
import media
import courseinfo
import supplies
from antibiotics import antibiotics
import chemicals

# --- Funtion to start a new experiment ---
def start_new_experiment():
    global current_experiment, biologicals_list, supplies_list, media_list, chemicals_list

    if current_experiment:
        experiments[current_experiment] = {
            "biologicals": biologicals_list.copy(),
            "supplies": supplies_list.copy(),
            "media": media_list.copy(),
            "chemicals": chemicals_list.copy()
        }

    print("Current experiment name:", current_experiment)
    print("Experiments stored:", experiments)
    
    # Set the current experiment to a new name (you may already have a method to get this)
    experiment_name = experiment_name_entry.get()  # or use another variable
    current_experiment = experiment_name  # Ensure it's assigned here

    # Initialize the lists for the new experiment
    
    biologicals_list = []
    supplies_list = []
    media_list = []
    chemicals_list = []

    print(f"Started new experiment: {current_experiment}")

# --- Functions for updating preview and adding/removing items ---
def update_preview():
    preview_text.config(state="normal")

    preview_text.insert(tk.END, f"{experiment_name_entry.get()}\n")

    preview_text.insert(tk.END, f"=== Biologicals ===\n")
    for b in biologicals_list:
        preview_text.insert(tk.END, f"{b['specimen']} on {b['media']} ({b['type']}) [{b['distribution']}]: ${b['cost']}\n")

    preview_text.insert(tk.END, f"\n=== Supplies ===\n")
    for s in supplies_list:
        preview_text.insert(tk.END, f"{s['name']} [{s['distribution']}]: ${s['cost']}\n")

    preview_text.insert(tk.END, f"\n=== Uninoculated Media ===\n")
    for m in media_list:
        preview_text.insert(tk.END, f"{m['media']} ({m['type']}) [{m['distribution']}]: ${m['cost']}\n")

    preview_text.insert(tk.END, f"\n=== Chemicals ===\n")
    for c in chemicals_list:
        preview_text.insert(tk.END, f"{c['chemical']} ({c['type']}) [{c['distribution']}]: ${c['cost']}\n")
    
    total_bio_cost = sum(b['cost'] for b in biologicals_list)
    total_supply_cost = sum(s['cost'] for s in supplies_list)
    total_media_cost = sum(m['cost'] for m in media_list)
    total_chemical_cost = sum(c['cost'] for c in chemicals_list)

    preview_text.insert(tk.END, f"\nTotal Biological Cost: ${round(total_bio_cost, 2)}\n")
    preview_text.insert(tk.END, f"Total Supplies Cost: ${round(total_supply_cost, 2)}\n")
    preview_text.insert(tk.END, f"Total Media Cost: ${round(total_media_cost, 2)}\n")
    preview_text.insert(tk.END, f"Total Chemical Cost: ${round(total_chemical_cost, 2)}\n")
    preview_text.insert(tk.END, f"Grand Total: ${round(total_bio_cost + total_supply_cost + total_media_cost + total_chemical_cost, 2)}\n\n")

    preview_text.config(state="disabled")

# --- Functions to add and remove biologicals, supplies, uninoculated media, and chemicals ---
def remove_last_biological():
    if biologicals_list:
        removed = biologicals_list.pop()
        print(f"Removed biological: {removed['specimen']}")
        update_preview()

def remove_last_supply():
    if supplies_list:
        removed = supplies_list.pop()
        print(f"Removed supply: {removed['name']}")
        update_preview()

def remove_last_uninoculated_media():
    if media_list:
        removed = media_list.pop()
        print(f"Removed uninoculated media: {removed['media']}")
        update_preview()

def remove_last_chemical():
    if chemicals_list:
        removed = chemicals_list.pop()
        print(f"Removed chemical: {removed['chemical']}")
        update_preview()

def add_biological():
    global biologicals_list

    selected_specimen = specimen_var.get()
    selected_media = media_var.get()
    selected_type = type_var.get()
    dist_num = distribution_entry.get()
    dist_type = distribution_type_var.get()

    if not all([selected_specimen, selected_media, selected_type, dist_num, dist_type]):
        print("Missing info for biological")
        return

    try:
        dist_num = int(dist_num)
    except ValueError:
        print("Distribution quantity must be a number")
        return

    course = courseinfo.courses.get(course_type_var.get())
    if not course:
        print("No course selected or invalid")
        return

    volume_ml = media.standard_volumes_ml[selected_type]
    media_data = next((v for k, v in media.media_list.items() if v["name"] == selected_media), None)

    if not media_data:
        print(f"No media data found for {selected_media}")
        return

    cost_per_ml = media_data["cost_per_ml"]
    total_samples = {
        "Per Student": course.students,
        "Per Pair": course.students // 2,
        "Per Group": course.groups,
        "Per Table": 6 * course.sections,
        "Per Section": course.sections,
        "Per Room": len(course.rooms),
        "Per Course": 1,
    }.get(dist_type, 1) * dist_num

    total_cost = volume_ml * cost_per_ml * total_samples

    biological = {
        "specimen": selected_specimen,
        "media": selected_media,
        "type": selected_type,
        "distribution": f"{dist_num} {dist_type}",
        "cost": round(total_cost, 2)
    }

    biologicals_list.append(biological)

    print(biologicals_list)

    update_preview()

def add_supply():
    global supplies_list

    selected_item_name = item_var.get()  # Get the selected supply name
    quantity = supply_quantity_entry.get()
    quantity_type = supply_quantity_type_var.get()

    if not all([selected_item_name, quantity, quantity_type]):
        print("Missing supply info")
        return

    try:
        quantity = int(quantity)
    except ValueError:
        print("Quantity must be a number")
        return

    # Get the course info
    course = courseinfo.courses.get(course_type_var.get())
    if not course:
        print("Invalid course selection")
        return

    # Determine if the selected item is a supply or antibiotic
    supply_key = supply_name_to_key.get(selected_item_name)
    antibiotic_key = antibiotic_name_to_key.get(selected_item_name)

    supply_info = supplies.supplies.get(supply_key)
    antibiotic_info = antibiotics.get(antibiotic_key)


    # If not found in supplies, check in antibiotics
    if not supply_info and not antibiotic_info:
        print("Item not found")
        return

    # If item is found in supplies
    if supply_info:
        per_item_cost = supply_info["cost_per_unit"] / supply_info["quantity"]
        category = "supply"
    # If item is found in antibiotics
    elif antibiotic_info:
        per_item_cost = antibiotic_info["cost_per_unit"] / antibiotic_info["quantity"]
        category = "antibiotic"

    # Calculate units based on quantity type
    units = {
        "Per Student": course.students,
        "Per Pair": course.students // 2,
        "Per Group": course.groups,
        "Per Table": 6 * course.sections,
        "Per Section": course.sections,
        "Per Room": len(course.rooms),
        "Per Course": 1,
    }.get(quantity_type, 1) * quantity

    total_cost = units * per_item_cost

    supplies_list.append({
        "name": selected_item_name,
        "category": category,  # Add category for reference (supply or antibiotic)
        "distribution": f"{quantity} {quantity_type}",
        "cost": round(total_cost, 2)
    })

    update_preview()

def add_uninoculated_media():
    global media_list

    selected_media = media_uninoc_var.get()
    selected_type = media_uninoc_type_var.get()
    dist_num = media_uninoc_quantity_entry.get()
    dist_type = media_uninoc_quantity_type_var.get()

    if not all([selected_media, selected_type, dist_num, dist_type]):
        print("Missing info for media")
        return

    try:
        dist_num = int(dist_num)
    except ValueError:
        print("Distribution quantity must be a number")
        return

    course = courseinfo.courses.get(course_type_var.get())
    if not course:
        print("No course selected or invalid")
        return

    volume_ml = media.standard_volumes_ml[selected_type]
    media_data = next((v for k, v in media.media_list.items() if v["name"] == selected_media), None)

    if not media_data:
        print(f"No media data found for {selected_media}")
        return

    cost_per_ml = media_data["cost_per_ml"]
    total_samples = {
        "Per Student": course.students,
        "Per Pair": course.students // 2,
        "Per Group": course.groups,
        "Per Table": 6 * course.sections,
        "Per Section": course.sections,
        "Per Room": len(course.rooms),
        "Per Course": 1,
    }.get(dist_type, 1) * dist_num

    total_cost = volume_ml * cost_per_ml * total_samples

    media_list.append({
        "media": selected_media,
        "type": selected_type,
        "distribution": f"{dist_num} {dist_type}",
        "cost": round(total_cost, 2)
    })

    update_preview()

def add_chemical():
    global chemicals_list

    selected_chemical = chemical_var.get()
    selected_type = chemical_type_var.get()
    dist_num = chemical_quantity_entry.get()
    dist_type = chemical_quantity_type_var.get()

    if not all([selected_chemical, selected_type, dist_num, dist_type]):
        print("Missing info for chemical")
        return

    try:
        dist_num = int(dist_num)
    except ValueError:
        print("Distribution quantity must be a number")
        return

    course = courseinfo.courses.get(course_type_var.get())
    if not course:
        print("No course selected or invalid")
        return
    
    for key, value in chemicals.chemical_list.items():
        if value.get("name") == selected_chemical:
            chemical_key = key
            break

    volume_ml = media.standard_volumes_ml[selected_type]
    chemical_data = chemicals.chemical_list.get(chemical_key)

    if not chemical_data:
        print(f"No chemical data found for {selected_chemical}")
        return

    cost_per_ml = chemical_data["cost_per_unit"]/chemical_data["quantity"]
    total_samples = {
        "Per Student": course.students,
        "Per Pair": course.students // 2,
        "Per Group": course.groups,
        "Per Table": 6 * course.sections,
        "Per Section": course.sections,
        "Per Room": len(course.rooms),
        "Per Course": 1,
    }.get(dist_type, 1) * dist_num

    total_cost = volume_ml * cost_per_ml * total_samples

    chemicals_list.append({
        "chemical": selected_chemical,
        "type": selected_type,
        "distribution": f"{dist_num} {dist_type}",
        "cost": round(total_cost, 2)
    })

    update_preview()

# --- GUI Layout Setup ---
# Store the position of the frame when it is initially placed
frame_position = {}

def toggle_frame(content_frame, toggle_button, frame_name):
    # Check if the content frame is currently visible
    if content_frame.winfo_ismapped():
        content_frame.grid_forget()  # Hide the content frame
        toggle_button.config(text="Show")  # Change button text to 'Show'
    else:
        # Use stored position from frame_position
        position = frame_position[frame_name]
        content_frame.grid(row=position['row'], column=position['column'],
                           columnspan=position['columnspan'], padx=10, pady=5, sticky="ew")  # Show the content frame
        toggle_button.config(text="Hide")  # Change button text to 'Hide'

root = tk.Tk()
root.title("Microbiology Lab Cost Calculator")
root.geometry("1200x900")

# Layout Frames
main_frame = ttk.Frame(root)
main_frame.grid(row=0, column=0, sticky="nsew")

input_frame = ttk.Frame(main_frame)
input_frame.grid(row=0, column=0, sticky="nw")

preview_frame = ttk.LabelFrame(main_frame, text="Preview")
preview_frame.grid(row=0, column=1, sticky="ne", padx=10, pady=10)

# Course Info
course_frame = ttk.LabelFrame(input_frame, text="Course Info")
course_frame.grid(row=0, column=0, columnspan=2, padx=10, pady=5, sticky="ew")

ttk.Label(course_frame, text="Course Name").grid(row=0, column=0, padx=5, pady=5)
course_type_var = tk.StringVar()
course_type_dropdown = ttk.Combobox(course_frame, textvariable=course_type_var, values=list(courseinfo.courses.keys()))
course_type_dropdown.grid(row=0, column=1, padx=5, pady=5)

# Experiment Name (Label + Entry + Set Button using grid)
ttk.Label(input_frame, text="Experiment Name:").grid(row=1, column=0, padx=5, pady=5, sticky="w")

experiment_name_entry = ttk.Entry(input_frame, width=40)
experiment_name_entry.grid(row=1, column=1, padx=5, pady=5)

start_new_button = ttk.Button(input_frame, text="Set Experiment Name", command=start_new_experiment)
start_new_button.grid(row=2, column=1, padx=5, pady=5)

# Biologicals Section
bio_frame = ttk.LabelFrame(input_frame, text="Biologicals")
bio_frame.grid(row=3, column=0, columnspan=2, padx=10, pady=5, sticky="ew")

bio_header_frame = ttk.Frame(bio_frame)
bio_header_frame.grid(row=0, column=0, columnspan=2, sticky="ew")

bio_toggle_button = ttk.Button(bio_header_frame, text="Hide", 
                                command=lambda: toggle_frame(bio_content_frame, bio_toggle_button, 'bio_content_frame'))
bio_toggle_button.grid(row=0, column=1, padx=5, pady=5)

bio_content_frame = ttk.Frame(bio_frame)
bio_content_frame.grid(row=2, column=0, columnspan=2, padx=10, pady=5, sticky="ew")

# Store the initial position of bio_content_frame
frame_position['bio_content_frame'] = {'row': 2, 'column': 0, 'columnspan': 2}

specimen_var = tk.StringVar()
specimen_dropdown = ttk.Combobox(bio_content_frame, textvariable=specimen_var, state="readonly", width=30)
specimen_dropdown.grid(row=1, column=1, padx=5, pady=5)
ttk.Label(bio_content_frame, text="Specimen:").grid(row=1, column=0, padx=5, pady=5)

media_var = tk.StringVar()
media_dropdown = ttk.Combobox(bio_content_frame, textvariable=media_var, state="readonly", width=40)
media_dropdown.grid(row=2, column=1, padx=5, pady=5)
ttk.Label(bio_content_frame, text="Media:").grid(row=2, column=0, padx=5, pady=5)

type_var = tk.StringVar()
type_dropdown = ttk.Combobox(bio_content_frame, textvariable=type_var, state="readonly", width=25)
type_dropdown.grid(row=3, column=1, padx=5, pady=5)
ttk.Label(bio_content_frame, text="Type:").grid(row=3, column=0, padx=5, pady=5)

distribution_entry = ttk.Entry(bio_content_frame, width=10)
distribution_entry.grid(row=4, column=1, padx=5, pady=5)
distribution_type_var = tk.StringVar()
distribution_type_dropdown = ttk.Combobox(bio_content_frame, textvariable=distribution_type_var, values=["Per Course", "Per Room", "Per Section", "Per Table", "Per Group", "Per Pair", "Per Student"])
distribution_type_dropdown.grid(row=4, column=2, padx=5, pady=5)
ttk.Label(bio_content_frame, text="Distribution:").grid(row=3, column=0, padx=5, pady=5)

media_dropdown.bind("<<ComboboxSelected>>", lambda event: type_dropdown.config(values=list(media.standard_volumes_ml.keys())))

add_biological_button = ttk.Button(bio_content_frame, text="Add Biological", command=add_biological)
add_biological_button.grid(row=5, column=0, columnspan=3, pady=5)

remove_bio_button = ttk.Button(bio_content_frame, text="Remove Last Biological", command=remove_last_biological)
remove_bio_button.grid(row=5, column=2, columnspan=3, pady=5)

# Supplies Section
supplies_frame = ttk.LabelFrame(input_frame, text="Supplies")
supplies_frame.grid(row=4, column=0, columnspan=2, padx=10, pady=5, sticky="ew")

supplies_header_frame = ttk.Frame(supplies_frame)
supplies_header_frame.grid(row=0, column=0, columnspan=2, sticky="ew")

supplies_toggle_button = ttk.Button(supplies_header_frame, text="Hide", 
                                command=lambda: toggle_frame(supplies_content_frame, supplies_toggle_button, 'supplies_content_frame'))
supplies_toggle_button.grid(row=0, column=1, padx=5, pady=5)

supplies_content_frame = ttk.Frame(supplies_frame)
supplies_content_frame.grid(row=2, column=0, columnspan=2, padx=10, pady=5, sticky="ew")

# Store the initial position of supplies_content_frame
frame_position['supplies_content_frame'] = {'row': 4, 'column': 0, 'columnspan': 2}

supply_names = [item["name"] for item in supplies.supplies.values()]
supply_name_to_key = {item["name"]: key for key, item in supplies.supplies.items()}

antibiotic_names = [item["name"] for item in antibiotics.values()]
antibiotic_name_to_key = {item["name"]: key for key, item in antibiotics.items()}

combined_items = supply_names + antibiotic_names

item_var = tk.StringVar()
item_dropdown = ttk.Combobox(supplies_content_frame, textvariable=item_var, values=combined_items, state="readonly", width=40)
item_dropdown.grid(row=0, column=1, padx=5, pady=5)
ttk.Label(supplies_content_frame, text="Supply:").grid(row=0, column=0, padx=5, pady=5)

supply_quantity_entry = ttk.Entry(supplies_content_frame, width=10)
supply_quantity_entry.grid(row=1, column=1, padx=5, pady=5)
ttk.Label(supplies_content_frame, text="Quantity:").grid(row=1, column=0, padx=5, pady=5)

supply_quantity_type_var = tk.StringVar()
quantity_type_dropdown = ttk.Combobox(supplies_content_frame, textvariable=supply_quantity_type_var, values=["Per Course", "Per Room", "Per Section", "Per Table", "Per Group", "Per Pair", "Per Student"])
quantity_type_dropdown.grid(row=1, column=3, padx=5, pady=5)
ttk.Label(supplies_content_frame, text="Quantity Type:").grid(row=1, column=2, padx=5, pady=5)

add_supply_button = ttk.Button(supplies_content_frame, text="Add Supply", command=add_supply)
add_supply_button.grid(row=2, column=0, columnspan=4, pady=5)

remove_supply_button = ttk.Button(supplies_content_frame, text="Remove Last Supply", command=remove_last_supply)
remove_supply_button.grid(row=2, column=2, columnspan=4, pady=5)

# Uninoculated Media Section
media_uninoc_frame = ttk.LabelFrame(input_frame, text="Uninoculated Media")
media_uninoc_frame.grid(row=5, column=0, columnspan=2, padx=10, pady=5, sticky="ew")

media_header_frame = ttk.Frame(media_uninoc_frame)
media_header_frame.grid(row=0, column=0, columnspan=2, sticky="ew")

media_toggle_button = ttk.Button(media_header_frame, text="Hide", 
                                command=lambda: toggle_frame(media_content_frame, media_toggle_button, 'media_content_frame'))
media_toggle_button.grid(row=0, column=1, padx=5, pady=5)

media_content_frame = ttk.Frame(media_uninoc_frame)
media_content_frame.grid(row=2, column=0, columnspan=2, padx=10, pady=5, sticky="ew")

# Store the initial position of media_content_frame
frame_position['media_content_frame'] = {'row': 6, 'column': 0, 'columnspan': 2}

media_uninoc_var = tk.StringVar()
media_uninoc_dropdown = ttk.Combobox(media_content_frame, textvariable=media_uninoc_var, state="readonly", width=40)
media_uninoc_dropdown.grid(row=0, column=1, padx=5, pady=5)
ttk.Label(media_content_frame, text="Media:").grid(row=0, column=0, padx=5, pady=5)

media_uninoc_type_var = tk.StringVar()
media_uninoc_type_dropdown = ttk.Combobox(media_content_frame, textvariable=media_uninoc_type_var, state="readonly", width=25)
media_uninoc_type_dropdown.grid(row=1, column=1, padx=5, pady=5)
ttk.Label(media_content_frame, text="Type:").grid(row=1, column=0, padx=5, pady=5)

media_uninoc_quantity_entry = ttk.Entry(media_content_frame, width=10)
media_uninoc_quantity_entry.grid(row=2, column=1, padx=5, pady=5)
media_uninoc_quantity_type_var = tk.StringVar()
media_uninoc_quantity_type_dropdown = ttk.Combobox(media_content_frame, textvariable=media_uninoc_quantity_type_var, values=["Per Course", "Per Room", "Per Section", "Per Table", "Per Group", "Per Pair", "Per Student"])
media_uninoc_quantity_type_dropdown.grid(row=2, column=2, padx=5, pady=5)
ttk.Label(media_content_frame, text="Distribution:").grid(row=2, column=0, padx=5, pady=5)

media_uninoc_dropdown.bind("<<ComboboxSelected>>", lambda event: media_uninoc_type_dropdown.config(values=list(media.standard_volumes_ml.keys())))
add_media_button = ttk.Button(media_content_frame, text="Add Media", command=add_uninoculated_media)
add_media_button.grid(row=3, column=0, columnspan=3, pady=5)

remove_media_button = ttk.Button(media_content_frame, text="Remove Last Media", command=remove_last_uninoculated_media)
remove_media_button.grid(row=3, column=2, columnspan=3, pady=5)

# Chemicals Section
chemicals_frame = ttk.LabelFrame(input_frame, text="Chemicals")
chemicals_frame.grid(row=6, column=0, columnspan=2, padx=10, pady=5, sticky="ew")

chemicals_header_frame = ttk.Frame(chemicals_frame)
chemicals_header_frame.grid(row=0, column=0, columnspan=2, sticky="ew")

chemicals_toggle_button = ttk.Button(chemicals_header_frame, text="Hide", 
                                command=lambda: toggle_frame(chemicals_content_frame, chemicals_toggle_button, 'chemicals_content_frame'))
chemicals_toggle_button.grid(row=0, column=1, padx=5, pady=5)

chemicals_content_frame = ttk.Frame(chemicals_frame)
chemicals_content_frame.grid(row=2, column=0, columnspan=2, padx=10, pady=5, sticky="ew")

# Store the initial position of chemical_content_frame
frame_position['chemicals_content_frame'] = {'row': 8, 'column': 0, 'columnspan': 2}

chemical_names = [item["name"] for item in chemicals.chemical_list.values()]

chemical_var = tk.StringVar()
chemical_dropdown = ttk.Combobox(chemicals_content_frame, textvariable=chemical_var, values=chemical_names, state="readonly", width=40)
chemical_dropdown.grid(row=0, column=1, padx=5, pady=5)
ttk.Label(chemicals_content_frame, text="Chemical:").grid(row=0, column=0, padx=5, pady=5)

chemical_type_var = tk.StringVar()
chemical_type_dropdown = ttk.Combobox(chemicals_content_frame, textvariable=chemical_type_var, state="readonly", width=25)
chemical_type_dropdown.grid(row=1, column=1, padx=5, pady=5)
ttk.Label(chemicals_content_frame, text="Type:").grid(row=1, column=0, padx=5, pady=5)

chemical_quantity_entry = ttk.Entry(chemicals_content_frame, width=10)
chemical_quantity_entry.grid(row=2, column=1, padx=5, pady=5)
chemical_quantity_type_var = tk.StringVar()
chemical_quantity_type_dropdown = ttk.Combobox(chemicals_content_frame, textvariable=chemical_quantity_type_var, values=["Per Course", "Per Room", "Per Section", "Per Table", "Per Group", "Per Pair", "Per Student"])
chemical_quantity_type_dropdown.grid(row=2, column=2, padx=5, pady=5)
ttk.Label(chemicals_content_frame, text="Distribution:").grid(row=2, column=0, padx=5, pady=5)

chemical_dropdown.bind("<<ComboboxSelected>>", lambda event: chemical_type_dropdown.config(values=list(media.standard_volumes_ml.keys())))
add_chemical_button = ttk.Button(chemicals_content_frame, text="Add Chemical", command=add_chemical)
add_chemical_button.grid(row=3, column=0, columnspan=3, pady=5)

remove_chemical_button = ttk.Button(chemicals_content_frame, text="Remove Last Chemical", command=remove_last_chemical)
remove_chemical_button.grid(row=4, column=0, columnspan=3, pady=5)

# Preview Text
preview_text = tk.Text(preview_frame, width=60, height=35, state="disabled", wrap="word")
preview_text.pack(padx=5, pady=5)

# Populate Dropdowns
if hasattr(bacteria, "bacteria_list"):
    specimen_dropdown["values"] = list(bacteria.bacteria_list)
if hasattr(media, "media_list"):
    media_dropdown["values"] = [value["name"] for key, value in media.media_list.items()]
if hasattr(media, "media_list"):
    media_uninoc_dropdown["values"] = [value["name"] for key, value in media.media_list.items()]

def add_italicized_text(para, text, italic_list):
    """
    Adds text to a paragraph, italicizing any matches from italic_list.
    """
    idx = 0
    while idx < len(text):
        match = None
        for name in italic_list:
            if text[idx:].startswith(name):
                match = name
                break

        if match:
            italic_run = para.add_run(match)
            italic_run.italic = True
            idx += len(match)
        else:
            para.add_run(text[idx])
            idx += 1

def write_biologicals(doc, biologicals):
    doc.add_heading("Biologicals", level=2)
    italic_list = list(bacteria.bacteria_list)

    for item in biologicals:
        para = doc.add_paragraph()

        specimen = item.get("specimen", "")
        media = item.get("media", "")
        type_ = item.get("type", "")
        distribution = item.get("distribution", "")
        cost = item.get("cost", "")

        # Don't show "premade" as type if it's already in media name
        if 'premade' in media:
            type_text = ""
        else:
            type_text = f"[{type_}] "

        add_italicized_text(para, specimen, italic_list)
        para.add_run(f" on {media} {type_text}[{distribution}]: ${cost}")

    doc.add_paragraph("")  # Blank line after section


def write_supplies(doc, supplies):
    doc.add_heading("Supplies", level=2)

    for item in supplies:
        para = doc.add_paragraph()
        name = item.get("name", "")
        distribution = item.get("distribution", "")
        cost = item.get("cost", "")

        para.add_run(f"Name: {name}, Distribution: {distribution}, Cost: ${cost}")

    doc.add_paragraph("")


def write_media(doc, media_items):
    doc.add_heading("Uninoculated Media", level=2)

    for item in media_items:
        para = doc.add_paragraph()
        media_name = item.get("media", "")
        type_ = item.get("type", "")
        distribution = item.get("distribution", "")
        cost = item.get("cost", "")

        para.add_run(f"Media: {media_name}, Type: {type_}, Distribution: {distribution}, Cost: ${cost}")

    doc.add_paragraph("")


def write_chemicals(doc, chemicals):
    doc.add_heading("Chemicals", level=2)

    for item in chemicals:
        para = doc.add_paragraph()
        chemical = item.get("chemical", "")
        type_ = item.get("type", "")
        distribution = item.get("distribution", "")
        cost = item.get("cost", "")

        para.add_run(f"Chemical: {chemical}, Type: {type_}, Distribution: {distribution}, Cost: ${cost}")

    doc.add_paragraph("")


# --- Main Export Function ---
def export_summary():
    global biologicals_list, supplies_list, media_list, chemicals_list, experiments

    if current_experiment:
        experiments[current_experiment] = {
            "biologicals": biologicals_list.copy(),
            "supplies": supplies_list.copy(),
            "media": media_list.copy(),
            "chemicals": chemicals_list.copy()
        }

    print("Current experiment name:", current_experiment)
    print("Experiments stored:", experiments)

    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word Documents", "*.docx")],
        title="{course_type_var.get()} Cost Analysis"
    )

    if not file_path:
        return

    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Consolas'  # or any font you want
    font.size = Pt(11)   # or any size you prefer
    
    # Title
    title = doc.add_heading(level=1)
    title_run = title.add_run(f"{course_type_var.get()} Cost Analysis n/ Based on Spring 2025")
    title_run.font.size = Pt()

    # Course Info
    course = courseinfo.courses.get(course_type_var.get())
    if course:
        course_paragraph = doc.add_paragraph()
        course_paragraph.add_run(f"Sections: {course.sections}  |  ")
        course_paragraph.add_run(f"Groups: {course.groups}  |  ")
        course_paragraph.add_run(f"Students: {course.students}  |  ")
        course_paragraph.add_run(f"Rooms: {course.rooms}")

    for experiment_name, data in experiments.items():
        print(data)

        # Big bold heading for experiment name
        exp_heading = doc.add_paragraph()
        exp_run = exp_heading.add_run(f"Experiment Name: {experiment_name}")
        exp_run.bold = True
        exp_run.underline = True
        exp_run.font.size = Pt(12)
        doc.add_paragraph("")

        # Write each section
        write_biologicals(doc, data["biologicals"])
        write_supplies(doc, data["supplies"])
        write_media(doc, data["media"])
        write_chemicals(doc, data["chemicals"])

        total_cost = data.get("total_cost", 0)  # Adjust if your structure differs
        total_paragraph = doc.add_paragraph()
        total_run = total_paragraph.add_run(f"Total Cost: ${total_cost:,.2f}")
        total_run.font.size = Pt(12)
        total_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    doc.save(file_path)
    print("DOCX Exported Successfully!")

export_button = ttk.Button(preview_frame, text="Export Summary", command=export_summary)
export_button.pack(pady=10)
   
root.mainloop()