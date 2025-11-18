import sys
import pandas as pd
import glob
import os
from docx import Document
from docx.shared import Pt
from datetime import datetime

# ----------------------------
# Setup paths and imports
# ----------------------------
script_dir = os.path.dirname(os.path.abspath(__file__))
library_path = os.path.abspath(os.path.join(script_dir, '..', 'Libraries'))
if library_path not in sys.path:
    sys.path.append(library_path)

import bacteria
import media
import supplies
import antibiotics
import chemicals

special_forms = ["1_ml", "0.5_ml", "1_ul"]

# ----------------------------
# COURSE + EXPERIMENT LOADING
# ----------------------------
def load_course_info(course_df):
    return {
        "students": int(course_df.loc[0, "Students"]),
        "sections": int(course_df.loc[0, "Sections"]),
        "groups": int(course_df.loc[0, "Groups"]),
        "rooms": course_df.loc[0, "Rooms"].split(",")
    }

def gather_all_items():
    folder_path = r"C:\Users\Kandria's PC\Desktop\ExpFiles"
    pattern = os.path.join(folder_path, "*Experiments.xlsx")
    experiment_files = glob.glob(pattern)

    if not experiment_files:
        raise FileNotFoundError("No *Experiments.xlsx files found.")

    # Assume only one file per course
    file = experiment_files[0]
    course_name = os.path.basename(file).split("Experiments.xlsx")[0].strip("_")
    xls = pd.ExcelFile(file)

    # Load course info
    course_df = xls.parse("CourseInfo")
    course_info = load_course_info(course_df)

    # Load experiment index
    experiment_list_df = xls.parse("ExperimentIndex")
    experiment_list_df.columns = experiment_list_df.columns.str.strip()

    # Load each experiment sheet using SheetName from index
    experiments = {}
    for _, row in experiment_list_df.iterrows():
        exp_title = str(row["Experiment Title"])
        sheet_name = str(row["SheetName"])
        if sheet_name.endswith('.0'):
            sheet_name = sheet_name[:-2]

        if sheet_name not in xls.sheet_names:
            print(f"WARNING: Sheet '{sheet_name}' missing — skipping.")
            continue

        exp_df = xls.parse(sheet_name)
        exp_df.columns = exp_df.columns.str.strip()
        exp_df.sheet_name = sheet_name  # attach sheet_name to df for later use
        experiments[exp_title] = exp_df

    return course_name, course_info, experiments

# ----------------------------
# LIBRARY LOOKUP
# ----------------------------
def lookup_by_name_or_key(name, library_dict, lib_label="item", sheet_name=None, exp_title=None):
    if not name or str(name).lower() == "nan":
        return None
    name_lower = str(name).strip().lower()

    for key, info in library_dict.items():
        if key.strip().lower() == name_lower:
            return info
        if str(info.get("name", "")).strip().lower() == name_lower:
            return info
        synonyms = info.get("synonyms") or info.get("synonym")
        if synonyms:
            if isinstance(synonyms, str):
                synonyms = [synonyms]
            if any(str(s).strip().lower() == name_lower for s in synonyms):
                return info

    context = ""
    if sheet_name:
        context += f" in sheet {sheet_name}"
    if exp_title:
        context += f" Experiment: '{exp_title}')"
    print(f"Warning: '{name}' not found in {lib_label} library{context}")
    return None

def validate_item(name, category, sheet_name=None, exp_title=None):
    if category == "Biological":
        return lookup_by_name_or_key(name, bacteria.bacteria_list, "bacteria", sheet_name, exp_title)
    elif category == "Uninoculated Media":
        return lookup_by_name_or_key(name, media.media_list, "media", sheet_name, exp_title)
    elif category == "Chemical":
        return lookup_by_name_or_key(name, chemicals.chemical_list, "chemical", sheet_name, exp_title)
    elif category == "Supply":
        return lookup_by_name_or_key(name, supplies.supplies, "supply", sheet_name, exp_title)
    elif category == "Antibiotics":
        return lookup_by_name_or_key(name, antibiotics.antibiotics, "antibiotic", sheet_name, exp_title)
    return None

# ----------------------------
# COST + QUANTITY CALCULATIONS
# ----------------------------
def calculate_total_qty(quantity, dist_type, course_info, form=None):
    multiplier = {
        "Per Student": course_info["students"],
        "Per Group": course_info["groups"],
        "Per Section": course_info["sections"],
        "Per Table": sum(5 if room.strip() == "113" else 6 for room in course_info['rooms']),
        "Per Room": len(course_info["rooms"]),
        "Per Course": 1
    }.get(dist_type, 1)
    return quantity * multiplier

def get_cost_per_use(info_dict):
    if info_dict is None:
        return 0

    cost = info_dict.get("cost_per_unit")
    if cost is None:
        cost = info_dict.get("cost_per_ml")

    if cost is None:
        print("Warning: No cost field in:", info_dict)
        return 0

    quantity = info_dict.get("quantity", 1)
    if quantity in [0, None]:
        quantity = 1
    try:
        return float(cost) / float(quantity)
    except Exception as e:
        print("Warning: Invalid numbers in:", info_dict, "Error:", e)
        return 0

def calculate_item_cost(row, course_info, exp_title=None, sheet_name=None):
    category = str(row.get("Category", "")).strip()
    name = str(row.get("Item/Organism", "")).strip()

    if category == "Biological":
        media_name = str(row.get("Media Used", "")).strip()
        if not media_name:
            return 0
        item_info = validate_item(media_name, "Uninoculated Media", sheet_name, exp_title)
    else:
        if pd.isna(name) and category == "Uninoculated Media":
            name = str(row.get("Media Used", "")).strip()
        item_info = validate_item(name, category, sheet_name, exp_title)

    if item_info is None:
        return 0

    cost_per_unit = get_cost_per_use(item_info)
    quantity = float(row.get("Quantity", 0))
    dist_type = str(row.get("DistributionType", "")).title()
    form = row.get("Form", None)

    total_qty_needed = calculate_total_qty(quantity, dist_type, course_info, form)
    return total_qty_needed * cost_per_unit

def calculate_experiment_cost(experiment_df, course_info, exp_title, sheet_name):
    total_cost = 0
    for _, row in experiment_df.iterrows():
        try:
            total_cost += calculate_item_cost(row, course_info, exp_title, sheet_name)
        except Exception as e:
            print(f"Error calculating cost for row: {e}")
    return total_cost

def calculate_course_cost(experiments_dict, course_info):
    output = {}
    grand_total = 0
    for exp_title, exp_df in experiments_dict.items():
        sheet_name = getattr(exp_df, 'sheet_name', exp_title)
        cost = calculate_experiment_cost(exp_df, course_info, exp_title, sheet_name)
        output[exp_title] = cost
        grand_total += cost
    output["Course Cost"] = grand_total
    return output

# ----------------------------
# WORD DOCUMENT EXPORT
# ----------------------------
def build_title(course_name):
    date_str = datetime.today().strftime("%B %Y")
    return f"{course_name} Cost Analysis — {date_str}"

def export_cost_assessment(course_name, course_info, experiments_dict, cost_dict):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Consolas'
    style.font.size = Pt(11)

    # Title
    doc.add_heading(build_title(course_name), level=1)

    # Course info
    doc.add_heading("Course Information", level=2)
    doc.add_paragraph(f"Students: {course_info['students']}")
    doc.add_paragraph(f"Sections: {course_info['sections']}")
    doc.add_paragraph(f"Groups: {course_info['groups']}")
    doc.add_paragraph(f"Rooms: {', '.join(course_info['rooms'])}")

    # Experiment costs
    doc.add_heading("Experiment Costs", level=2)
    for exp_title, cost in cost_dict.items():
        if exp_title != "Course Cost":
            doc.add_paragraph(f"{exp_title}: ${cost:.2f}")

    # Total Cost
    doc.add_heading("Total Course Cost", level=2)
    doc.add_paragraph(f"${cost_dict['Course Cost']:.2f}")

    doc.save(f"{course_name}_Cost_Analysis.docx")

# ----------------------------
# MAIN EXECUTION
# ----------------------------
if __name__ == "__main__":
    course_name, course_info, experiments = gather_all_items()
    cost_dict = calculate_course_cost(experiments, course_info)
    export_cost_assessment(course_name, course_info, experiments, cost_dict)
    print("Word document created!")


# import sys
# import pandas as pd
# import glob
# import os
# from openpyxl import load_workbook
# from docx import Document
# from docx.shared import Pt
# from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
# from datetime import datetime

# script_dir = os.path.dirname(os.path.abspath(__file__))
# library_path = os.path.abspath(os.path.join(script_dir, '..', 'Libraries'))
# if library_path not in sys.path:
#     sys.path.append(library_path)

# import bacteria
# import media
# import supplies
# import antibiotics
# import chemicals

# special_forms = ["1_ml", "0.5_ml", "1_ul"]

# def load_course_info(course_df):
#     return {
#         "students":int(course_df.loc[0, "Students"]),
#         "sections":int(course_df.loc[0, "Sections"]),
#         "groups":int(course_df.loc[0, "Groups"]),
#         "rooms": course_df.loc[0, "Rooms"].split(",")
#     }

# def gather_all_items(course_name=None): ##doesnt have cost stuff in it yet, just pulling from doc 

#     folder_path = r"C:\Kandriad\Desktop\ExpFiles"
#     pattern = os.path.join(folder_path, "*Experiments.xlsx")
#     experiment_files = glob.glob(pattern)

#     for file in experiment_files:
#         course_name = os.path.basename(file).split("Experiments.xlsx")[0]
#         xls = pd.ExcelFile(file)
#         course_df = xls.parse("CourseInfo")
#         course_info = load_course_info(course_df)

#     experiment_list_df = xls.parse("ExperimentIndex")
#     experiment_list_df.columns = experiment_list_df.columns.str.strip()

#     for _, row in experiment_list_df.iterrows():
#         exp_title = row["Experiment Title"]
#         sheet_name = str(row["SheetName"])
#         if sheet_name.endswith(".0"):
#             sheet_name = sheet_name[:-2]

#         if sheet_name not in xls.sheet_names:
#             print(f"Sheet '{sheet_name}' not found in file '{file}' -- Skipping.")
#             continue

#     experiment_df = xls.parse(experiment_list_df)
#     experiment_df.columns = experiment_df.columns.str.strip()

#     for _, item in experiment_df.iterrows():
#         item_category = str(item.get("Category", "")).stip()
#         if not item_category or item_category.lower() == "nan":
#             continue

#         item_name = item.get("Item/Organism")

#         if pd.isna(item_name) and item_category == "Uninoculated Media":
#             item_name == item.get("Media Used")
#             name = str(item_name).strip() if not pd.isna(item_name) else ""
#             media_used = item.get("Media Used", 0)
#             form = item.get("Form", 0)
#             quantity = item.get("Quantity", 0)
#             dist_type = str(item.get("DistributionType", "")).strip().title()

#             try:
#                 total_qty = calculate_total_qty(quantity, dist_type, course_info, form)
#             except Exception as e:
#                 print(f"Error calculating quantity for {course_name} {exp_title} - {item_category} {name} {media_used}: {e}")
#                 continue
#     return course_name, course_info, experiment_df

# def lookup_by_name_or_key(name, library_dict, lib_label="item"):
#     if not name or str(name).lower() == "nan":
#         return None
#     name_lower = str(name).strip().lower()

#     for key, info in library_dict.items():
#         if key.strip().lower() == name_lower:
#             return info
        
#         if str(info.get("name", "")).strip().lower() == name_lower:
#             return info
        
#         synonyms = info.get("synonym" or "synonyms", [])
#         if any(syn.strip().lower == name_lower for syn in synonyms):
#             return info
        
#     print(f"Warning: '{name}' not found in {lib_label} library")
#     return None

# def is_valid_bacteria(name):
#     return lookup_by_name_or_key(name, bacteria.bacteria_list, "bacteria") is not None
# def is_valid_media(name):
#     return lookup_by_name_or_key(name, media.media_list, "media") is not None
# def is_valid_supply(name):
#     return lookup_by_name_or_key(name, supplies.supplies, "supply") is not None
# def is_valid_chemical(name):
#     return lookup_by_name_or_key(name, chemicals.chemical_list, "chemical") is not None
# def is_valid_antibiotic(name):
#     return lookup_by_name_or_key(name, antibiotics.antibiotics, "antibiotic") is not None

# def validate_item(name, category):
#     if category == "Biological":
#         return is_valid_bacteria(name)
#     elif category == "Uninoculated Media":
#         return is_valid_media(name)
#     elif category == "Chemical":
#         return is_valid_chemical(name)
#     elif category == "Supply":
#         return is_valid_supply(name)
#     elif category == "Antibiotics":
#         return is_valid_antibiotic(name)

# def calculate_total_qty(quantity, dist_type, course_info, form=None):
#     multiplier = {
#         "Per Student": course_info["students"],
#         "Per Group": course_info["groups"],\
#         "Per Section": course_info["sections"],
#         "Per Table": sum(5 if room.strip() == "113" else 6 for room in course_info['rooms']),
#         "Per Room": len(course_info["rooms"]),
#         "Per Course": 1
#     }.get(dist_type, 1)
    
#     return quantity * multiplier #special forms shouldnt matter here because the end goal is cost, not how many samples to prepare

# def get_cost_per_use(info_dict):
#     if info_dict is None:
#         return 0
    
#     qty = info_dict.get("quantity")
#     if qty is None or   ty == 0:
#         print("Warning: No valid quanitity field in:", info_dict)
#         return 0
    
#     cpu = info_dict.get("cost_per_unit")
#     if cpu is (None, 0):
#         print("Warning: No valid cost per unit field in:", info_dict)
#         return 0
    
#     try:
#         return float(cpu) / float(qty)
#     except:
#         print("Warning: invalid numbers:", info_dict)
#         return 0

# def calculate_item_cost(row, course_info):
#     category = str(row.get("Category", "")).strip()
#     name = str(row.get("Item/Organism", "")).strip()

#     if pd.isna(name) and category == "Uninoculated Media":
#         name = str(row.get("Media Used", "")).strip() 

#     item_info = validate_item(name, category)
#     if item_info is None:
#         return 0
    
#     cost_per_use = get_cost_per_use(item_info)

#     quantity = float(row.get("Quantity", 0))
#     dist_type = str(row.get("DistributionType", "")).title()
#     form = row.get("Form", None) 

#     total_qty_needed = calculate_total_qty(quantity, dist_type, course_info, form)  

#     return total_qty_needed * cost_per_use

# def calculate_experiment_cost(experiment_df, course_info):
#     exp_cost = 0
    
#     for _, row in experiment_df.iterrows():
#         try:
#             exp_cost += calculate_item_cost(row, course_info)
#         except Exception as e:
#             print(f"Error processing item {row.get('Item/Organism')}: {e}")
#     return exp_cost

# def calculate_course_cost(all_exp_dict, course_info):
#     result = {}
#     grand_total = 0

#     for exp_title, exp_df in all_exp_dict.items():
#         cost = calculate_experiment_cost(exp_df, course_info)
#         result[exp_title] = cost
#         grand_total += cost

#     result["Course Cost"] = grand_total
#     return result

# ##create a word doc with Exp Costs and Course Costs to present
# def build_title(course_name):
#     date_str = datetime.today().strftime("%B %Y")
#     title = f"{course_name} Cost Analysis {date_str}"
#     return title

# def export_cost_assessment(course_name, experiment_df, cost_data):
#     doc = Document()

#     style = doc.styles['Normal']
#     font = style.font
#     font.name = 'Consolas'
#     font.size = Pt(11)

#     title = build_title(course_name)
#     heading1 = doc.add_heading(title, level=1)
#     course_info = doc.add_heading(load_course_info, level=2)
#     course_cost = f"{calculate_course_cost}"
#     exp_with_cost = f"{experiment_df}: {calculate_experiment_cost}"

#     doc.save(f"{course_name}_Cost_Analysis.docx")